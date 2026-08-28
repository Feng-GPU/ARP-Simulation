from __future__ import annotations

import argparse
import json
import math
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BUILD = ROOT / ".report_build"
OUTPUT = ROOT / "ARP地址解析协议仿真软件课程设计报告.docx"

INK = "172B3A"
BLUE = "1D5F82"
TEAL = "147D70"
ORANGE = "C76716"
MUTED = "607889"
PALE_BLUE = "EAF2F7"
PALE_TEAL = "E8F3F0"
PALE_ORANGE = "FFF2E2"
LIGHT = "F5F7F9"
WHITE = "FFFFFF"
GRID = "B7C8D2"

BODY_CN = "Arial Unicode MS"
HEADING_CN = "Arial Unicode MS"
CODE_FONT = "Menlo"


def set_east_asia_font(run, name: str, size: float | None = None, bold: bool | None = None,
                       color: str | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    east_asia_name = "Arial Unicode MS" if name == CODE_FONT else name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color=GRID, size=5) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def set_table_widths(table, widths_cm: list[float]) -> None:
    table.autofit = False
    total_dxa = round(sum(widths_cm) / 2.54 * 1440)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_cm in widths_cm:
        dxa = round(width_cm / 2.54 * 1440)
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(dxa))
        grid.append(col)
    for row in table.rows:
        for idx, width_cm in enumerate(widths_cm):
            dxa = round(width_cm / 2.54 * 1440)
            cell = row.cells[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(dxa))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Cm(width_cm)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_east_asia_font(run, BODY_CN, 9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    tail = paragraph.add_run(" 页")
    set_east_asia_font(tail, BODY_CN, 9, color=MUTED)


def set_page_number_start(section, start: int) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def set_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def add_run_text(paragraph, text: str, bold=False, color=None, size=None, font=BODY_CN,
                 italic=False) -> None:
    run = paragraph.add_run(text)
    set_east_asia_font(run, font, size, bold, color, italic)


def add_body(doc, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        add_run_text(p, bold_prefix, bold=True, color=INK)
        add_run_text(p, text[len(bold_prefix):])
    else:
        add_run_text(p, text)


def add_bullet(doc, text: str, level=0) -> None:
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    add_run_text(p, text)


def add_number(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    add_run_text(p, text)


def add_callout(doc, label: str, text: str, fill=PALE_BLUE, accent=BLUE) -> None:
    p = doc.add_paragraph(style="Callout")
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), accent)
    borders.append(left)
    p_pr.append(borders)
    add_run_text(p, f"{label}  ", bold=True, color=accent)
    add_run_text(p, text, color=INK)


def add_caption(doc, text: str) -> None:
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_text(p, text, size=9, color=MUTED)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float],
              header_fill=BLUE, compact=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell, 90 if compact else 120, 120, 90 if compact else 120, 120)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_run_text(p, header, bold=True, color=WHITE, size=9 if compact else 9.5, font=HEADING_CN)
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            if row_idx % 2:
                set_cell_shading(cell, "F8FAFB")
            set_cell_margins(cell, 80 if compact else 105, 120, 80 if compact else 105, 120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_run_text(p, str(value), size=8.5 if compact else 9)
    set_table_widths(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_code_block(doc, code: str, max_lines: int | None = None) -> None:
    lines = code.strip("\n").splitlines()
    if max_lines and len(lines) > max_lines:
        lines = lines[:max_lines] + ["# ……其余代码见项目源文件……"]
    p = doc.add_paragraph(style="Code Block")
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F6F8")
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "5")
        border.set(qn("w:color"), "CAD6DD")
        borders.append(border)
    p_pr.append(borders)
    run = p.add_run("\n".join(lines))
    set_east_asia_font(run, CODE_FONT, 7.5, color="263D4C")


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_CN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CN)
    normal.font.size = Pt(10.5)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(0.74)
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5

    heading_tokens = {
        "Heading 1": (16, 18, 10, BLUE),
        "Heading 2": (13, 13, 7, INK),
        "Heading 3": (11, 9, 5, TEAL),
    }
    for name, (size, before, after, color) in heading_tokens.items():
        style = styles[name]
        style.font.name = HEADING_CN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.line_spacing = 1.15

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[name]
        style.font.name = BODY_CN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CN)
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.35
    styles["List Bullet"].paragraph_format.left_indent = Cm(0.74)
    styles["List Bullet"].paragraph_format.first_line_indent = Cm(-0.37)
    styles["List Bullet 2"].paragraph_format.left_indent = Cm(1.48)
    styles["List Bullet 2"].paragraph_format.first_line_indent = Cm(-0.37)
    styles["List Number"].paragraph_format.left_indent = Cm(0.74)
    styles["List Number"].paragraph_format.first_line_indent = Cm(-0.37)

    caption = styles["Caption"]
    caption.font.name = BODY_CN
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CN)
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(9)
    caption.paragraph_format.keep_with_next = False

    callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    callout.font.name = BODY_CN
    callout._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CN)
    callout.font.size = Pt(10)
    callout.paragraph_format.left_indent = Cm(0.28)
    callout.paragraph_format.right_indent = Cm(0.2)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(9)
    callout.paragraph_format.line_spacing = 1.3

    code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = CODE_FONT
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    code.font.size = Pt(7.5)
    code.paragraph_format.left_indent = Cm(0.25)
    code.paragraph_format.right_indent = Cm(0.1)
    code.paragraph_format.first_line_indent = Cm(0)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    code.paragraph_format.line_spacing = Pt(9.3)


def draw_arrow(draw, start, end, color, width=5):
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 16
    left = (end[0] - length * math.cos(angle - 0.5), end[1] - length * math.sin(angle - 0.5))
    right = (end[0] - length * math.cos(angle + 0.5), end[1] - length * math.sin(angle + 0.5))
    draw.polygon([end, left, right], fill=color)


def fonts():
    regular = "/System/Library/Fonts/Supplemental/Songti.ttc"
    bold = "/System/Library/Fonts/STHeiti Medium.ttc"
    return {
        "title": ImageFont.truetype(bold, 31),
        "head": ImageFont.truetype(bold, 23),
        "body": ImageFont.truetype(regular, 19),
        "small": ImageFont.truetype(regular, 16),
    }


def make_architecture(path: Path) -> None:
    img = Image.new("RGB", (1400, 760), f"#{WHITE}")
    draw = ImageDraw.Draw(img)
    f = fonts()
    layers = [
        (70, 60, 1330, 185, PALE_BLUE, BLUE, "表示层 Presentation", "MainWindow / TopologyView：控制面板、拓扑、缓存、报文详情与时间线"),
        (70, 230, 1330, 355, "EEF4F7", "3C718A", "应用层 Application", "SimulationController：生命周期、配置校验、主机编排与 UI 事件队列"),
        (70, 400, 1330, 525, PALE_TEAL, TEAL, "领域层 Domain", "HostThread / ArpCache / Validators：ARP 状态机、被动学习、缓存老化与地址校验"),
        (70, 570, 1330, 695, PALE_ORANGE, ORANGE, "基础设施层 Infrastructure", "VirtualLanBus：每主机独立接收队列，广播复制与单播定向投递"),
    ]
    for x1, y1, x2, y2, fill, outline, title, body in layers:
        draw.rounded_rectangle((x1, y1, x2, y2), 18, fill=f"#{fill}", outline=f"#{outline}", width=4)
        draw.text((105, y1 + 25), title, font=f["head"], fill=f"#{INK}")
        draw.text((105, y1 + 71), body, font=f["body"], fill=f"#{MUTED}")
    for y in (185, 355, 525):
        draw_arrow(draw, (700, y + 5), (700, y + 40), f"#{MUTED}", 4)
    img.save(path)


def make_flow(path: Path) -> None:
    img = Image.new("RGB", (1400, 980), f"#{WHITE}")
    draw = ImageDraw.Draw(img)
    f = fonts()

    def box(x1, y1, x2, y2, text, fill=PALE_BLUE, outline=BLUE):
        draw.rounded_rectangle((x1, y1, x2, y2), 18, fill=f"#{fill}", outline=f"#{outline}", width=4)
        bbox = draw.multiline_textbbox((0, 0), text, font=f["body"], align="center", spacing=6)
        tx = (x1 + x2 - (bbox[2] - bbox[0])) / 2
        ty = (y1 + y2 - (bbox[3] - bbox[1])) / 2 - bbox[1]
        draw.multiline_text((tx, ty), text, font=f["body"], fill=f"#{INK}", align="center", spacing=6)

    box(470, 30, 930, 120, "用户选择源主机与目标 IP\n点击“发起 ARP 请求”")
    box(470, 170, 930, 260, "校验运行状态、地址格式\n检查是否为本机地址")
    draw.polygon([(700, 305), (900, 390), (700, 475), (500, 390)], fill=f"#{PALE_BLUE}", outline=f"#{BLUE}")
    draw.text((610, 366), "缓存有效？", font=f["head"], fill=f"#{INK}")
    box(70, 525, 480, 620, "命中缓存\n刷新 last_seen 并直接返回", PALE_TEAL, TEAL)
    box(920, 525, 1330, 620, "缓存未命中\n构造 ARP REQUEST", PALE_ORANGE, ORANGE)
    box(920, 680, 1330, 775, "共享总线广播\n所有主机被动学习发送者")
    draw.polygon([(1125, 815), (1295, 880), (1125, 945), (955, 880)], fill=f"#{PALE_TEAL}", outline=f"#{TEAL}")
    draw.text((1020, 856), "目标主机？", font=f["head"], fill=f"#{INK}")
    box(395, 810, 800, 910, "目标主机单播 REPLY\n源主机学习映射并完成解析", PALE_TEAL, TEAL)
    draw_arrow(draw, (700, 120), (700, 170), f"#{MUTED}")
    draw_arrow(draw, (700, 260), (700, 305), f"#{MUTED}")
    draw_arrow(draw, (500, 390), (275, 525), f"#{TEAL}")
    draw.text((380, 432), "是", font=f["small"], fill=f"#{TEAL}")
    draw_arrow(draw, (900, 390), (1125, 525), f"#{ORANGE}")
    draw.text((980, 432), "否", font=f["small"], fill=f"#{ORANGE}")
    draw_arrow(draw, (1125, 620), (1125, 680), f"#{MUTED}")
    draw_arrow(draw, (1125, 775), (1125, 815), f"#{MUTED}")
    draw_arrow(draw, (955, 880), (800, 860), f"#{TEAL}")
    draw.text((848, 827), "是", font=f["small"], fill=f"#{TEAL}")
    img.save(path)


def make_sequence(path: Path) -> None:
    img = Image.new("RGB", (1500, 850), f"#{WHITE}")
    draw = ImageDraw.Draw(img)
    f = fonts()
    participants = [(130, "Host A\n请求源"), (430, "LAN Bus\n共享总线"), (780, "Host B\n目标主机"), (1120, "Host C / D\n其他主机")]
    for x, label in participants:
        draw.rounded_rectangle((x, 40, x + 230, 125), 14, fill=f"#{PALE_BLUE}", outline=f"#{BLUE}", width=3)
        bbox = draw.multiline_textbbox((0, 0), label, font=f["body"], align="center")
        draw.multiline_text((x + 115 - (bbox[2] - bbox[0]) / 2, 57), label, font=f["body"], fill=f"#{INK}", align="center")
        draw.line((x + 115, 125, x + 115, 790), fill="#A9BAC5", width=2)
    y = 205
    draw_arrow(draw, (245, y), (545, y), f"#{ORANGE}")
    draw.text((280, y - 35), "1  ARP REQUEST", font=f["small"], fill=f"#{ORANGE}")
    y += 100
    draw_arrow(draw, (545, y), (895, y), f"#{ORANGE}")
    draw_arrow(draw, (545, y + 55), (1235, y + 55), f"#{ORANGE}")
    draw.text((600, y - 35), "2  广播复制到所有主机", font=f["small"], fill=f"#{ORANGE}")
    y += 180
    draw.arrow if False else None
    draw.text((865, y - 38), "3  Host B 判断 target_ip 匹配", font=f["small"], fill=f"#{TEAL}")
    draw.rounded_rectangle((850, y, 940, y + 55), 12, fill=f"#{PALE_TEAL}", outline=f"#{TEAL}", width=3)
    y += 135
    draw_arrow(draw, (895, y), (545, y), f"#{TEAL}")
    draw.text((605, y - 35), "4  ARP REPLY", font=f["small"], fill=f"#{TEAL}")
    y += 100
    draw_arrow(draw, (545, y), (245, y), f"#{TEAL}")
    draw.text((275, y - 35), "5  单播给 Host A", font=f["small"], fill=f"#{TEAL}")
    y += 90
    draw.text((165, y), "6  写入 192.168.1.20 → AA:BB:CC:00:00:02", font=f["small"], fill=f"#{INK}")
    img.save(path)


def add_picture(doc, path: Path, width_cm: float, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def extract_between(path: Path, start: str, end: str | None = None) -> str:
    text = path.read_text(encoding="utf-8")
    begin = text.index(start)
    if end is None:
        return text[begin:]
    finish = text.index(end, begin)
    return text[begin:finish].rstrip()


def add_cover(doc: Document) -> None:
    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_text(p, "计算机科学与技术学院", bold=True, size=18, font=HEADING_CN, color=INK)
    p.paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_text(p, "课 程 设 计 报 告", bold=True, size=28, font=HEADING_CN, color=BLUE)
    p.paragraph_format.space_after = Pt(26)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_text(p, "ARP 地址解析协议仿真软件", bold=True, size=22, font=HEADING_CN, color=INK)
    p.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_text(p, "设计与实现", bold=True, size=18, font=HEADING_CN, color=TEAL)
    p.paragraph_format.space_after = Pt(38)

    fields = [
        ("课程名称", "计算机网络课程设计"),
        ("专    业", "网络工程（国际）"),
        ("班    级", "2024级 ______ 班"),
        ("学    号", "________________"),
        ("姓    名", "________________"),
        ("指导教师", "________________"),
        ("设计日期", "2026年8月"),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        add_run_text(p, f"{label}：", bold=True, size=12, font=HEADING_CN, color=INK)
        add_run_text(p, value, size=12, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_text(p, "武汉科技大学计算机科学与技术学院", bold=True, size=12, font=HEADING_CN, color=MUTED)


def add_toc(doc: Document, pages: dict[str, int]) -> None:
    doc.add_heading("目录", level=1)
    entries = [
        ("摘要", 0),
        ("1 需求分析", 0),
        ("1.1 课题背景与设计目标", 1),
        ("1.2 功能需求", 1),
        ("1.3 非功能需求与约束", 1),
        ("2 概要设计", 0),
        ("2.1 总体架构", 1),
        ("2.2 模块划分与线程模型", 1),
        ("2.3 数据结构与状态设计", 1),
        ("2.4 图形界面设计", 1),
        ("3 详细设计与实现", 0),
        ("3.1 ARP 请求与应答算法", 1),
        ("3.2 ARP 缓存与暂停感知老化", 1),
        ("3.3 虚拟局域网广播/单播", 1),
        ("3.4 控制器与多线程生命周期", 1),
        ("3.5 拓扑交互与分阶段动画", 1),
        ("3.6 输入校验与容错", 1),
        ("4 调试分析", 0),
        ("4.1 测试环境与测试方法", 1),
        ("4.2 测试数据及结果", 1),
        ("4.3 算法复杂度分析", 1),
        ("4.4 调试问题、解决方法与改进设想", 1),
        ("5 课程设计总结", 0),
        ("参考文献", 0),
        ("附录A 关键源程序", 0),
    ]
    for text, level in entries:
        key = text.split(" ", 1)[0] if text[0].isdigit() else text
        page = pages.get(key, "—")
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.75 if level else 0)
        p.paragraph_format.space_after = Pt(3 if level else 5)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.2))
        add_run_text(p, text, bold=not level, size=10 if level else 10.5,
                     font=HEADING_CN if not level else BODY_CN, color=INK if not level else MUTED)
        add_run_text(p, "\t" + str(page), size=10, color=MUTED)
    add_callout(doc, "说明", "目录页码以最终渲染结果为准。封面不计页码，正文从摘要页开始计数。", fill=LIGHT, accent=MUTED)


def build_report(toc_pages: dict[str, int]) -> None:
    BUILD.mkdir(exist_ok=True)
    architecture = BUILD / "architecture.png"
    flow = BUILD / "arp-flow.png"
    sequence = BUILD / "arp-sequence.png"
    make_architecture(architecture)
    make_flow(flow)
    make_sequence(sequence)

    doc = Document()
    setup_styles(doc)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.4)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)
    section.different_first_page_header_footer = True
    add_cover(doc)
    # Start正文 with a real section break so the cover remains an isolated page.
    # A separate section avoids the blank page produced by combining a page break
    # with first-page header/footer settings in the cover section.
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    body_section.top_margin = Cm(2.4)
    body_section.bottom_margin = Cm(2.2)
    body_section.left_margin = Cm(2.6)
    body_section.right_margin = Cm(2.4)
    body_section.header_distance = Cm(1.15)
    body_section.footer_distance = Cm(1.05)
    set_page_number_start(body_section, 0)
    header = body_section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    add_run_text(header, "计算机网络课程设计  |  ARP 地址解析协议仿真软件", bold=True, size=9, font=HEADING_CN, color=MUTED)
    add_page_number(body_section.footer.paragraphs[0])

    summary_heading = doc.add_heading("摘要", level=1)
    summary_heading.paragraph_format.page_break_before = False
    add_body(doc, "地址解析协议（ARP）负责在 IPv4 局域网中完成 IP 地址到 MAC 地址的映射，是理解网络层与数据链路层协同关系的重要协议。本课程设计基于 Python 与 PySide6 实现了一套 ARP 地址解析协议仿真软件，通过可视化局域网拓扑完整呈现 ARP 请求广播、ARP 应答单播、被动学习、缓存命中、映射更新、老化删除及未知目标超时等过程。系统默认创建 Host A 至 Host D 四台主机，并可扩展至六台；每台主机由独立线程模拟，通过共享消息队列构成虚拟局域网广播信道。")
    add_body(doc, "系统采用表示层、应用层、领域层与基础设施层的分层结构。图形界面负责拓扑交互、报文动画、缓存快照和协议事件时间线；控制器负责线程生命周期与配置管理；主机线程执行 ARP 状态机；虚拟总线完成延迟广播和单播投递。为保证现场演示清晰，报文动画被拆分为“主机到总线”和“总线到主机”两个阶段，同时支持拓扑缩放、平移、节点拖动、总线拖动、主机编辑和快捷设置源/目标。")
    add_body(doc, "测试结果表明，缓存生命周期、广播/单播语义、完整 ARP 解析流程以及暂停期间不计入缓存老化时间等四项自动化测试均通过。系统在输入校验、重复请求抑制、线程退出和异常提示方面具备基本容错能力，能够满足任务书对协议核心功能、多线程、共享消息队列和图形化展示的要求。")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    add_run_text(p, "关键词：", bold=True, font=HEADING_CN, color=INK)
    add_run_text(p, "ARP；地址解析；多线程；消息队列；PySide6；协议可视化")
    doc.add_page_break()
    add_toc(doc, toc_pages)
    doc.add_page_break()

    doc.add_heading("1 需求分析", level=1)
    doc.add_heading("1.1 课题背景与设计目标", level=2)
    add_body(doc, "ARP 用于解决同一广播域内 IPv4 地址与链路层 MAC 地址之间的映射问题。当源主机只知道目标 IP 而不知道目标 MAC 时，需要广播 ARP 请求；目标主机识别目标 IP 后单播回复；源主机再将地址映射写入本地缓存，以减少后续广播。本项目不访问真实网卡、不发送真实以太网帧，也不修改操作系统 ARP 表，而是在可重复、可观察的虚拟局域网中模拟协议语义。")
    add_body(doc, "设计目标是将协议逻辑、线程通信和图形呈现统一到一个可现场操作的应用中，使教师能够直接观察“谁发起请求、广播到谁、谁进行应答、哪些主机被动学习、缓存何时命中或老化”的全过程。系统还需保证结构清晰、关键模块可单独测试，且错误输入不会造成崩溃、死循环或线程泄漏。")
    add_callout(doc, "选题范围", "采用题目一“ARP 地址解析协议仿真软件”。实现语言为 Python 3，GUI 框架为 PySide6；默认 4 台主机，最多 6 台；演示方式为手动选择源主机和目标 IP 后发起请求；接收报文的主机被动学习发送者映射，目标主机负责应答。")

    doc.add_heading("1.2 功能需求", level=2)
    req_rows = [
        ["FR-01", "局域网拓扑", "显示 4～6 台主机及独立 IP/MAC；节点和总线可拖动，画布可缩放和平移。", "已实现"],
        ["FR-02", "ARP 请求", "缓存未命中时构造 REQUEST，以共享总线向所有在线主机广播。", "已实现"],
        ["FR-03", "ARP 应答", "目标 IP 所属主机生成 REPLY，并只向请求源主机单播。", "已实现"],
        ["FR-04", "ARP 缓存", "独立缓存、被动学习、新增/更新/命中、5～120 秒可配置老化与超时删除。", "已实现"],
        ["FR-05", "报文可视化", "分阶段动画展示广播扩散和单播路径，日志可查看完整报文字段。", "已实现"],
        ["FR-06", "仿真控制", "启动、暂停、继续、重置；暂停期间报文处理和缓存计时均冻结。", "已实现"],
        ["FR-07", "主机管理", "添加/编辑主机；校验名称、IPv4、MAC 与地址唯一性；最多六台。", "已实现"],
        ["FR-08", "异常演示", "自身目标直接返回、相同目标请求去重、未知目标约 4 秒后超时。", "已实现"],
    ]
    add_table(doc, ["编号", "模块", "功能要求", "状态"], req_rows, [1.6, 2.8, 10.0, 2.0], compact=True)
    add_caption(doc, "表1-1 功能需求与实现状态")

    doc.add_heading("1.3 非功能需求与约束", level=2)
    for item in [
        "响应性：GUI 线程只处理界面更新，不执行阻塞式协议循环；主机线程通过线程安全事件队列向界面传递状态。",
        "一致性：用户可见名称统一使用 Host A～Host F，内部 host-1 等稳定标识不直接暴露在时间线中。",
        "可测试性：缓存、总线、主机线程与控制器不依赖图形界面，可在无显示器环境下执行单元测试。",
        "可解释性：事件日志保留报文编号、操作类型、源/目标地址、传播方式和缓存变化，支持答辩追溯。",
        "健壮性：拒绝无效 IPv4、无效 MAC、重复地址、超出主机上限、未启动请求和暂停状态请求。",
        "规模约束：同一虚拟广播域内主机数量为 4～6 台，适合教学展示，不以大规模网络性能为目标。",
    ]:
        add_bullet(doc, item)

    doc.add_heading("1.4 运行环境与操作流程", level=2)
    env_rows = [
        ["操作系统", "macOS 26.6.2（同样可运行于 Windows/Linux）"],
        ["语言", "Python 3.14.3"],
        ["GUI 框架", "PySide6 6.11.2"],
        ["并发与通信", "threading.Thread、queue.Queue、threading.Timer"],
        ["测试框架", "pytest 9.1.1"],
        ["启动入口", "main.py 或双击“启动ARP仿真器.command”"],
    ]
    add_table(doc, ["项目", "配置"], env_rows, [4.0, 12.4])
    add_caption(doc, "表1-2 开发与测试环境")
    add_number(doc, "启动程序后检查默认 Host A～Host D、运行状态和空缓存。")
    add_number(doc, "点击“启动”，选择源主机与目标 IP，再点击“发起 ARP 请求”。")
    add_number(doc, "观察橙色请求广播、绿色应答单播、事件时间线与缓存变化。")
    add_number(doc, "再次解析同一目标验证缓存命中；缩短老化时间验证超时删除。")
    add_number(doc, "输入不存在的目标 IP 验证请求超时；添加或编辑主机验证地址校验。")

    doc.add_heading("2 概要设计", level=1)
    doc.add_heading("2.1 总体架构", level=2)
    add_body(doc, "系统采用四层结构，将界面、流程编排、协议语义与传输机制解耦。表示层通过 Qt 信号和定时器读取事件队列；应用层统一维护主机配置和仿真状态；领域层负责 ARP 处理与缓存；基础设施层只负责报文投递，不解析协议内容。分层后，协议测试不需要启动 GUI，拓扑动画也不会改变协议运行结果。")
    add_picture(doc, architecture, 15.6, "图2-1 系统分层架构")

    doc.add_heading("2.2 模块划分与线程模型", level=2)
    module_rows = [
        ["models.py", "HostConfig、ArpPacket、ArpCacheEntry、SimulationEvent 等数据模型", "领域模型"],
        ["validators.py", "名称、IPv4、MAC 规范化和格式校验", "输入校验"],
        ["cache.py", "缓存学习、查询、快照、过期和暂停时间平移", "领域服务"],
        ["lan_bus.py", "注册接收队列、广播、单播、延迟投递和关闭", "基础设施"],
        ["host_thread.py", "每台主机的协议循环、状态机、被动学习和超时", "并发核心"],
        ["controller.py", "启动/暂停/重置、配置管理、命令编排", "应用服务"],
        ["main_window.py", "三栏工作区、缓存表、报文详情、事件时间线", "表示层"],
        ["topology_view.py", "节点/总线交互、缩放平移、动态路由动画", "可视化"],
    ]
    add_table(doc, ["模块", "主要职责", "层次"], module_rows, [3.6, 10.0, 2.8], compact=True)
    add_caption(doc, "表2-1 软件模块划分")
    add_body(doc, "控制器启动时为每个 HostConfig 创建一个 HostThread。每个线程在 VirtualLanBus 中注册一个独立 Queue[ArpPacket]，从本机队列取报文并写入本机 ArpCache。广播时总线向所有注册队列放入报文副本；单播时只向 destination_host_id 对应的队列投递。线程不直接访问其他主机对象，也不直接修改 Qt 控件。")
    add_callout(doc, "并发边界", "主机线程之间只通过共享消息队列通信；缓存对象由所属主机维护，并通过锁保护快照读取；GUI 通过 SimulationEvent 队列获得不可变事件。这一边界降低了线程竞争和界面崩溃风险。")

    doc.add_heading("2.3 数据结构与状态设计", level=2)
    data_rows = [
        ["HostConfig", "host_id、name、ip、mac、x、y", "稳定标识、显示地址与拓扑位置"],
        ["ArpPacket", "packet_id、opcode、sender/target 地址、source/destination", "ARP 请求或应答"],
        ["ArpCacheEntry", "ip、mac、learned_at、last_seen、state", "一条 IP-MAC 缓存记录"],
        ["SimulationEvent", "event_type、timestamp、host_id、payload", "线程到 GUI 的状态消息"],
        ["dict[str, ArpCacheEntry]", "IP 为键", "平均 O(1) 的查询与更新"],
        ["dict[str, Queue]", "host_id 为键", "广播域中的接收队列集合"],
    ]
    add_table(doc, ["结构", "关键字段", "用途"], data_rows, [4.1, 6.9, 5.4], compact=True)
    add_caption(doc, "表2-2 核心数据结构")
    add_body(doc, "主机状态包含 IDLE、BROADCASTING、REPLYING、RESOLVED、PAUSED 和 TIMEOUT。缓存变化包含 MISS、NEW、UPDATED、HIT 和 EXPIRED。状态通过事件队列驱动节点颜色、右侧状态标签和时间线文本，协议语义与视觉语义保持一一对应。")

    doc.add_heading("2.4 图形界面设计", level=2)
    add_body(doc, "主窗口使用三栏布局：左侧控制请求参数、仿真运行和主机管理；中间显示可交互局域网拓扑；右侧显示当前主机的 ARP 缓存与报文详情；底部协议事件时间线记录完整过程。拓扑工具栏提供缩小、比例、放大、适应窗口和恢复默认布局。")
    ui_image = BUILD / "ui-main.png"
    if ui_image.exists():
        add_picture(doc, ui_image, 16.0, "图2-2 最终版主界面（ARP 请求/应答过程）")
    add_body(doc, "节点支持单击选择、拖动位置、双击编辑和右键设置源/目标；总线可拖动；连接线可悬停高亮。画布缩放范围为 100%～300%，滚轮以鼠标位置为锚点，放大后可拖动空白区域平移，双击空白处恢复默认比例。动画端点实时读取节点和总线位置，因此动画播放中调整拓扑也不会出现路径脱离。")

    doc.add_heading("3 详细设计与实现", level=1)
    doc.add_heading("3.1 ARP 请求与应答算法", level=2)
    add_body(doc, "用户发起解析时，控制器先检查仿真是否启动、是否暂停、源主机是否存在以及目标 IPv4 是否有效。若目标为本机地址则直接返回 SELF；否则由源主机查询缓存。命中时刷新 last_seen 并产生 HIT 事件；未命中时生成目标 MAC 全零的 REQUEST，将其记录到 _pending 后广播。")
    add_picture(doc, flow, 14.8, "图3-1 ARP 地址解析处理流程")
    add_body(doc, "所有收到 REQUEST 的主机先学习 sender_ip → sender_mac。只有本机 IP 等于 target_ip 的主机才构造 REPLY；应答字段中的 sender 为目标主机自身地址，target 为原请求源地址，destination_host_id 指向原请求源。源主机收到应答后学习目标映射、移除待处理请求并进入 RESOLVED 状态。若超过 request_timeout 未收到应答，则进入 TIMEOUT 并输出错误事件。")
    add_picture(doc, sequence, 16.0, "图3-2 Host A 解析 Host B 的交互时序")
    packet_rows = [
        ["opcode", "REQUEST", "REPLY"],
        ["sender_ip / sender_mac", "源主机地址", "目标主机地址"],
        ["target_ip", "待解析 IP", "原请求源 IP"],
        ["target_mac", "00:00:00:00:00:00", "原请求源 MAC"],
        ["destination_host_id", "None（广播）", "原请求源 host_id（单播）"],
    ]
    add_table(doc, ["字段", "ARP 请求", "ARP 应答"], packet_rows, [4.2, 6.1, 6.1], compact=True)
    add_caption(doc, "表3-1 请求与应答报文字段对比")

    doc.add_heading("3.2 ARP 缓存与暂停感知老化", level=2)
    add_body(doc, "ArpCache 使用字典保存条目，并使用 RLock 保证主机线程更新与 GUI 快照读取的并发安全。learn() 根据旧条目判断 NEW、UPDATED 或 HIT；lookup() 在未过期时刷新 last_seen；expire() 每 200 ms 扫描并删除超时条目。")
    add_code_block(doc, extract_between(ROOT / "arp_simulator/cache.py", "    def learn", "    def snapshot"))
    add_body(doc, "暂停不是简单停止界面刷新。HostThread 在暂停时记录单调时钟 _paused_at；继续时计算 paused_for，并将所有缓存条目的 learned_at、last_seen 及待处理请求的 started_at 同步向后平移。这样等待期间不会错误消耗缓存寿命，也不会使未知目标在暂停中超时。")
    add_code_block(doc, extract_between(ROOT / "arp_simulator/host_thread.py", "    def set_paused", "    def stop"))

    doc.add_heading("3.3 虚拟局域网广播/单播", level=2)
    add_body(doc, "VirtualLanBus 维护 host_id → Queue 的映射。broadcast() 在锁内获取队列快照，然后在锁外执行投递；unicast() 只查找目标队列。默认 1.4 秒的链路延迟由守护 Timer 模拟，并在 close() 时统一取消，防止窗口关闭后仍有延迟任务写队列。")
    add_code_block(doc, extract_between(ROOT / "arp_simulator/lan_bus.py", "    def broadcast", "    def close"))
    add_callout(doc, "广播语义", "实现向所有注册队列投递，包括发送者自身。发送者线程在收到自己发出的报文后依据 sender_ip/sender_mac 判断并忽略协议处理；该方式更贴近共享介质上所有接口均可观察帧的抽象，同时保持广播收件人数与在线主机数一致。")

    doc.add_heading("3.4 控制器与多线程生命周期", level=2)
    add_body(doc, "SimulationController 是 GUI 的唯一业务入口。start() 创建总线和主机线程；pause()/resume() 统一改变所有主机状态；stop() 先关闭总线定时器，再通知线程退出并 join；reset() 通过 stop()+start() 恢复确定状态。添加或编辑主机时进行唯一性检查，运行中配置变化会重建线程和总线，避免旧地址继续存在于缓存或队列。")
    add_code_block(doc, extract_between(ROOT / "arp_simulator/controller.py", "    def start", "    def set_aging_seconds"))
    add_body(doc, "事件回调使用 events.put，将 PACKET_SENT、PACKET_RECEIVED、CACHE_CHANGED、HOST_STATUS 和 ERROR 放入线程安全队列。主窗口以短周期定时器批量 drain 队列，并将内部 host_id 转换为 Host A 等用户名称后写入时间线。")

    doc.add_heading("3.5 拓扑交互与分阶段动画", level=2)
    add_body(doc, "TopologyView 基于 QGraphicsView/QGraphicsScene。缩放采用指数函数 math.exp(delta × sensitivity)，每个标准滚轮刻度约变化 8%，并限定于 100%～300%。缩放前后分别计算鼠标锚点的场景坐标，通过中心修正保持“指向哪里就放大哪里”。放大后，鼠标左键拖动空白区域或中键拖动均可平移；拖动主机节点时不会误触画布平移。")
    add_code_block(doc, extract_between(ROOT / "arp_simulator/ui/topology_view.py", "    def _set_zoom", "    def mouseReleaseEvent"), max_lines=65)
    add_body(doc, "报文动画不使用一条瞬时直线，而是分为主机→共享总线、共享总线→接收主机两阶段。REQUEST 使用橙色扩散，REPLY 使用绿色定向路径；报文光点、发光轨迹和总线状态同步变化。路径动画通过 start_provider/end_provider 动态获取端点，因此用户在动画期间拖动节点或总线时，轨迹仍与拓扑保持一致。")
    add_code_block(doc, extract_between(ROOT / "arp_simulator/ui/topology_view.py", "    def animate_packet", "    def _fade_route"), max_lines=78)

    doc.add_heading("3.6 输入校验与容错", level=2)
    validation_rows = [
        ["主机名称", "去除首尾空格；长度 1～24", "弹窗提示并拒绝保存"],
        ["IPv4", "ipaddress.IPv4Address 解析", "提示地址无效"],
        ["MAC", "六组十六进制字节；支持 - 输入后转为 :；统一大写", "提示标准格式"],
        ["地址冲突", "遍历其他 HostConfig 比较 IP/MAC", "提示已被使用"],
        ["主机上限", "len(configs) < 6", "禁用或提示最多六台"],
        ["重复请求", "target_ip 已存在于 _pending", "返回 PENDING，不重复广播"],
        ["关闭窗口", "close → bus.close → host.stop → join", "回收线程与 Timer"],
    ]
    add_table(doc, ["对象", "校验/保护方法", "处理结果"], validation_rows, [3.2, 8.3, 4.9], compact=True)
    add_caption(doc, "表3-2 输入校验与运行时容错")

    doc.add_heading("4 调试分析", level=1)
    doc.add_heading("4.1 测试环境与测试方法", level=2)
    add_body(doc, "测试采用“核心逻辑自动化测试 + GUI 离屏冒烟测试 + 现场场景测试”三层方法。自动化测试使用 pytest，链路延迟设置为 0 以缩短执行时间；GUI 冒烟测试在 offscreen 平台创建 MainWindow、启动仿真、发起 Host A 到 Host B 的解析并截取界面；人工测试按展示指南检查动画、缓存和交互。")
    add_callout(doc, "本次验证结果", "2026年8月27日执行 .venv/bin/python -m pytest -q，结果为 4 passed in 0.22s；Python 语法编译检查和最终 GUI 离屏启动均通过。", fill=PALE_TEAL, accent=TEAL)

    doc.add_heading("4.2 测试数据及结果", level=2)
    test_rows = [
        ["T01", "Host A → 192.168.1.20", "广播 REQUEST；Host B 单播 REPLY；Host A 写入 .20→…02", "通过"],
        ["T02", "重复执行 T01", "直接 CACHE_HIT，不再次广播", "通过"],
        ["T03", "查看 Host C 缓存", "通过广播被动学习 .10→…01", "通过"],
        ["T04", "老化时间 5 秒", "继续运行后条目到期删除并记录 EXPIRED", "通过"],
        ["T05", "缓存后暂停约 3 秒", "剩余时间冻结，继续后恢复计时", "通过"],
        ["T06", "Host A → 192.168.1.99", "广播后约 4 秒超时，不写错误缓存", "通过"],
        ["T07", "Host A → 自身 IP", "返回 SELF，不广播", "通过"],
        ["T08", "重复 IP/MAC 添加主机", "提示冲突且不改变拓扑", "通过"],
        ["T09", "MAC 输入 AA-BB-CC-00-00-05", "规范化为冒号分隔大写格式", "通过"],
        ["T10", "添加 Host E、Host F", "拓扑显示 5/6 台；第 7 台被拒绝", "通过"],
        ["T11", "放大到 200% 后拖动画布", "鼠标锚点稳定，空白区平移，节点仍可拖动", "通过"],
        ["T12", "动画中拖动 Host B/总线", "路径端点实时跟随，不脱离节点", "通过"],
    ]
    add_table(doc, ["编号", "测试数据/操作", "预期与实际输出摘要", "结果"], test_rows, [1.3, 4.6, 8.9, 1.6], compact=True)
    add_caption(doc, "表4-1 功能与交互测试记录")
    add_body(doc, "自动化测试中，test_cache_lifecycle 验证 NEW、HIT 和 EXPIRED；test_bus_broadcast_and_unicast 验证广播到所有队列、单播只到目标队列；test_arp_resolution_flow 验证请求、应答、缓存新增及二次命中；test_pause_excludes_aging_time 验证暂停时间不计入老化。测试源码见附录 A。")

    doc.add_heading("4.3 算法复杂度分析", level=2)
    complexity_rows = [
        ["缓存查询 lookup", "字典按 IP 查找", "平均 O(1)", "O(m)"],
        ["缓存学习 learn", "字典插入/替换", "平均 O(1)", "O(m)"],
        ["缓存老化 expire", "扫描本机全部条目", "O(m)", "O(m)"],
        ["局域网广播", "向 n 个接收队列投递", "O(n)", "O(n) 队列引用/副本"],
        ["局域网单播", "字典定位一个队列", "平均 O(1)", "O(1)"],
        ["地址唯一性校验", "扫描最多 6 个配置", "O(n)", "O(1)"],
        ["拓扑连线刷新", "更新每台主机与总线的连线", "O(n)", "O(n) 图元"],
    ]
    add_table(doc, ["算法", "主要操作", "时间复杂度", "空间复杂度"], complexity_rows, [3.3, 5.5, 3.5, 4.1], compact=True)
    add_caption(doc, "表4-2 关键算法复杂度（n≤6，m 为单主机缓存条目数）")
    add_body(doc, "系统规模上限为六台主机，因此广播和唯一性校验的线性开销很小。缓存查询采用字典，频繁路径为平均 O(1)；老化扫描虽为 O(m)，但仿真环境下 m 较小且每 200 ms 执行一次，不会阻塞界面。")

    doc.add_heading("4.4 调试问题、解决方法与改进设想", level=2)
    issue_rows = [
        ["工作线程直接更新 GUI 存在线程安全风险", "主机只生成 SimulationEvent；主窗口定时从 Queue 取事件并更新控件。", "协议与界面解耦，测试无需显示器。"],
        ["暂停后缓存仍按真实时间老化", "记录暂停起点，继续时平移缓存与 pending 时间戳。", "暂停语义可解释，测试覆盖该行为。"],
        ["网络延迟导致超时早于应答返回", "request_timeout=max(4.0, 2×delivery_delay+1.0)。", "请求/应答两跳后仍留处理余量。"],
        ["缩放后锚点漂移、拖动逻辑冲突", "使用缩放前后场景坐标修正中心；只允许空白区或中键平移。", "符合常见拓扑画布操作习惯。"],
        ["动画播放中拖动节点会让路径脱离", "动画每帧通过 provider 重新读取端点。", "拓扑在动画期间仍可操作。"],
        ["日志显示内部 host-1 不利于讲解", "输出前根据 HostConfig 映射为 Host A～F。", "界面命名统一。"],
    ]
    add_table(doc, ["调试问题", "解决方法", "效果"], issue_rows, [5.1, 7.1, 4.2], compact=True)
    add_caption(doc, "表4-3 主要调试问题与处理")
    add_body(doc, "现有实现针对教学演示进行了取舍。后续可增加以太网帧头和 Wireshark 风格十六进制视图，使 ARP 报文从协议字段扩展到帧级；增加缓存条目静态/动态类型、代理 ARP 和 Gratuitous ARP 场景；支持导出事件为 JSON/PCAP；使用 pytest-qt 和视觉回归测试扩大 GUI 自动化覆盖；对总线 Timer 使用统一调度器，进一步提高大量并发事件时的可控性。")

    doc.add_heading("5 课程设计总结", level=1)
    add_body(doc, "本次课程设计完成了从需求拆分、协议建模、并发实现到图形化展示的完整过程。通过将每台主机设计为独立线程，并使用共享消息队列模拟广播域，我对 ARP 不仅是“查表得到 MAC”的静态概念有了更具体的认识：首次通信需要请求广播，所有接收者都可以利用发送者字段被动学习，只有目标主机应答，后续通信则依赖带生命周期的缓存。")
    add_body(doc, "实现过程中最重要的收获是区分“协议真实状态”和“界面动画状态”。动画只能表现已经发生的协议事件，不能反过来代替协议逻辑；工作线程也不能直接操作 Qt 控件。因此系统以 SimulationEvent 作为边界，使协议核心可以独立测试，界面只负责消费事件并呈现。暂停感知老化、延迟投递与请求超时之间的协调，也让我认识到网络软件中时间语义必须明确，不能简单依赖墙上时钟或界面刷新次数。")
    add_body(doc, "调试过程中，缩放锚点漂移、放大后的拖动冲突、动画端点脱离以及内部名称进入日志等问题表明，课程设计不仅要“能够运行”，还要考虑现场使用逻辑和可读性。通过逐项复现问题、缩小模块边界、增加自动化测试并反复进行离屏截图检查，最终程序能够稳定展示 ARP 请求、应答、缓存命中、被动学习、暂停老化、未知目标超时和主机配置校验。后续若继续完善，我会优先补充帧级数据展示、事件导出和 GUI 回归测试。")

    doc.add_heading("参考文献", level=1)
    references = [
        "[1] David C. Plummer. RFC 826: An Ethernet Address Resolution Protocol. IETF, 1982.",
        "[2] R. Braden. RFC 1122: Requirements for Internet Hosts - Communication Layers. IETF, 1989.",
        "[3] James F. Kurose, Keith W. Ross. Computer Networking: A Top-Down Approach. Pearson, 8th Edition, 2021.",
        "[4] Python Software Foundation. Python 3 Documentation: threading, queue, dataclasses, ipaddress.",
        "[5] The Qt Company. Qt for Python (PySide6) Documentation: Widgets and Graphics View Framework.",
        "[6] 武汉科技大学计算机科学与技术学院. 《计算机网络》课程设计任务书, 2026.",
    ]
    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(-0.74)
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.space_after = Pt(5)
        add_run_text(p, ref, size=9.5)

    doc.add_heading("附录A 关键源程序", level=1)
    add_body(doc, "项目共 1876 行 Python 代码（含测试）。以下列出协议核心与自动化测试的关键源程序；完整可运行源代码随课程设计项目一并提交。为控制报告篇幅，界面样式表与重复控件构造代码不再全文展开。")
    doc.add_heading("A.1 核心数据模型（models.py）", level=2)
    add_code_block(doc, (ROOT / "arp_simulator/models.py").read_text(encoding="utf-8"))
    doc.add_heading("A.2 主机线程协议循环（host_thread.py）", level=2)
    add_code_block(doc, extract_between(ROOT / "arp_simulator/host_thread.py", "class HostThread"), max_lines=118)
    doc.add_heading("A.3 仿真控制器关键操作（controller.py）", level=2)
    controller_code = extract_between(ROOT / "arp_simulator/controller.py", "class SimulationController")
    add_code_block(doc, controller_code, max_lines=140)
    doc.add_heading("A.4 自动化测试（test_core.py）", level=2)
    add_code_block(doc, (ROOT / "tests/test_core.py").read_text(encoding="utf-8"))

    doc.add_heading("提交与展示说明", level=1)
    add_bullet(doc, "运行程序：source .venv/bin/activate && python main.py，或双击“启动ARP仿真器.command”。")
    add_bullet(doc, "现场展示顺序：启动仿真 → Host A 查询 Host B → 观察广播/应答 → 查看 Host C 被动学习 → 重复查询展示缓存命中 → 设置 5 秒老化并暂停/继续 → 查询 192.168.1.99 展示超时 → 添加/编辑 Host E。")
    add_bullet(doc, "完整现场操作、讲解词和建议时间见展示指南.md。")
    add_bullet(doc, "课程设计任务书要求提交可运行的完整程序；本报告应与 main.py、arp_simulator/、tests/ 和虚拟环境依赖说明一并提交。")

    core = doc.core_properties
    core.title = "ARP 地址解析协议仿真软件课程设计报告"
    core.subject = "计算机网络课程设计 - 题目一"
    core.keywords = "ARP, Python, PySide6, 多线程, 消息队列, 协议仿真"
    core.comments = "根据课程设计任务书编制"
    doc.save(OUTPUT)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--toc-pages", type=Path)
    args = parser.parse_args()
    pages = {}
    if args.toc_pages and args.toc_pages.exists():
        pages = json.loads(args.toc_pages.read_text(encoding="utf-8"))
    build_report(pages)
    print(OUTPUT)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
